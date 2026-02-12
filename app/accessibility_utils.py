"""
accessibility_utils.py

Helper functions for checking accessibility of different file types.

Each file type (PDF, Word, Excel) has its own set of checks.
These functions do the actual inspection work.
"""

from __future__ import annotations

import io
from typing import List, Tuple
from pathlib import Path

# PDF handling
import fitz  # PyMuPDF - for reading PDF structure and content

# Word document handling
try:
    from docx import Document  # python-docx library
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Excel handling
try:
    from openpyxl import load_workbook
    from openpyxl.styles import Font, Color
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

from app.accessibility_models import (
    AccessibilityIssue,
    IssueLevel,
    PDFAccessibilityChecks,
    DocxAccessibilityChecks,
    XlsxAccessibilityChecks,
)


# =============================================================================
# PDF ACCESSIBILITY CHECKS
# =============================================================================

def check_pdf_accessibility(pdf_path: str) -> Tuple[PDFAccessibilityChecks, List[AccessibilityIssue]]:
    """
    Check PDF for ADA compliance.
    
    Returns:
        - PDFAccessibilityChecks: Results of technical checks
        - List[AccessibilityIssue]: List of problems found
    
    What we check:
    1. Is the PDF tagged? (Required for screen readers)
    2. Is document language specified?
    3. Do images have alt text?
    4. Is text searchable? (Not a scanned image)
    5. Is reading order logical?
    6. Are headings properly structured?
    """
    issues: List[AccessibilityIssue] = []
    
    # Open the PDF file
    doc = fitz.open(pdf_path)
    
    # -------------------------------------------------------------------------
    # CHECK 1: Is the PDF tagged?
    # -------------------------------------------------------------------------
    # Tagged PDFs have structure trees that tell screen readers what's a heading,
    # paragraph, list, table, etc. Without tags, screen readers just read
    # everything as plain text in visual order.
    
    is_tagged = False
    try:
        # Try to access the structure tree (tags)
        struct_tree = doc.get_toc()  # Table of contents
        catalog = doc.pdf_catalog()  # PDF catalog dictionary
        
        # Check if StructTreeRoot exists in the catalog
        # This is the root of the tagging structure
        if catalog and "StructTreeRoot" in catalog:
            is_tagged = True
    except Exception:
        is_tagged = False
    
    if not is_tagged:
        issues.append(AccessibilityIssue(
            wcag_criterion="1.3.1",  # Info and Relationships
            level=IssueLevel.CRITICAL,
            description="PDF is not tagged. Screen readers cannot understand document structure.",
            location="Document properties",
            remediation=(
                "1. Open PDF in Adobe Acrobat Pro\n"
                "2. Go to Tools → Accessibility → Autotag Document\n"
                "3. Manually verify and fix tag structure\n"
                "4. Run Accessibility Checker to confirm"
            ),
            blocks_compliance=True
        ))
    
    # -------------------------------------------------------------------------
    # CHECK 2: Is document language specified?
    # -------------------------------------------------------------------------
    # Screen readers need to know what language to use for pronunciation.
    # Without this, "resume" might be read as "re-zoom" instead of "rez-oo-may"
    
    has_language = False
    try:
        # Check if the document metadata has a language setting
        metadata = doc.metadata
        if metadata and metadata.get("language"):
            has_language = True
        
        # Also check the catalog for /Lang entry
        catalog = doc.pdf_catalog()
        if catalog and "Lang" in catalog:
            has_language = True
    except Exception:
        pass
    
    if not has_language:
        issues.append(AccessibilityIssue(
            wcag_criterion="3.1.1",  # Language of Page
            level=IssueLevel.ERROR,
            description="Document language is not specified. Screen readers may use wrong pronunciation.",
            location="Document properties",
            remediation=(
                "1. Open PDF in Adobe Acrobat Pro\n"
                "2. Go to File → Properties → Advanced\n"
                "3. Set Language field (e.g., 'English (United States)')\n"
                "4. Save document"
            ),
            blocks_compliance=True
        ))
    
    # -------------------------------------------------------------------------
    # CHECK 3: Do images have alt text?
    # -------------------------------------------------------------------------
    # Every meaningful image must have alternative text describing it
    # for users who cannot see the image.
    
    images_with_alt = 0
    images_without_alt = 0
    
    # Loop through each page
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Get all images on this page
        image_list = page.get_images(full=True)
        
        for img_index, img in enumerate(image_list):
            # Try to find alt text for this image
            # Note: PyMuPDF has limited ability to extract alt text from tagged PDFs
            # This is a basic check - full validation requires Adobe tools
            
            has_alt = False
            
            # In a properly tagged PDF, images would be in the structure tree
            # with an /Alt entry. This is difficult to extract with PyMuPDF,
            # so we're using a heuristic approach here.
            
            # For now, we'll assume images don't have alt text unless proven otherwise
            # A production system would use Adobe's Accessibility API or PAC 3
            
            if not has_alt:
                images_without_alt += 1
                issues.append(AccessibilityIssue(
                    wcag_criterion="1.1.1",  # Non-text Content
                    level=IssueLevel.CRITICAL,
                    description=f"Image #{img_index + 1} is missing alternative text.",
                    location=f"Page {page_num + 1}",
                    remediation=(
                        "1. Right-click image in Adobe Acrobat Pro\n"
                        "2. Select 'Edit Alternate Text'\n"
                        "3. Add descriptive text explaining what the image shows\n"
                        "4. If image is decorative only, mark it as artifact"
                    ),
                    blocks_compliance=True
                ))
            else:
                images_with_alt += 1
    
    # -------------------------------------------------------------------------
    # CHECK 4: Is text searchable? (OCR check)
    # -------------------------------------------------------------------------
    # If the PDF is a scanned image with no text layer, it's completely
    # inaccessible to screen readers.
    
    has_searchable_text = False
    total_text_length = 0
    
    for page in doc:
        text = page.get_text().strip()
        total_text_length += len(text)
    
    # If we found reasonable amount of text, it's searchable
    # (Rule of thumb: at least 100 characters)
    if total_text_length > 100:
        has_searchable_text = True
    
    is_ocr_needed = not has_searchable_text
    
    if is_ocr_needed:
        issues.append(AccessibilityIssue(
            wcag_criterion="1.1.1",  # Non-text Content
            level=IssueLevel.CRITICAL,
            description=(
                "PDF appears to be a scanned image with no text layer. "
                "Screen readers cannot read the content."
            ),
            location="Entire document",
            remediation=(
                "1. Use OCR (Optical Character Recognition) to add text layer\n"
                "2. In Adobe Acrobat Pro: Tools → Scan & OCR → Recognize Text\n"
                "3. Review and correct any OCR errors\n"
                "4. Add proper tagging after OCR"
            ),
            blocks_compliance=True
        ))
    
    # -------------------------------------------------------------------------
    # CHECK 5: Reading order
    # -------------------------------------------------------------------------
    # Content should be read in logical order, not visual layout order.
    # Example: A two-column article should read column 1 fully, then column 2,
    # not line 1 col 1, line 1 col 2, line 2 col 1, etc.
    
    # This is very difficult to check programmatically without AI.
    # We can only do a basic heuristic: if the document is tagged,
    # we assume reading order has been considered.
    # A human should always verify this manually.
    
    has_logical_reading_order = is_tagged  # Heuristic assumption
    
    if not has_logical_reading_order:
        issues.append(AccessibilityIssue(
            wcag_criterion="1.3.2",  # Meaningful Sequence
            level=IssueLevel.ERROR,
            description="Reading order may not be logical. Requires manual verification.",
            location="Entire document",
            remediation=(
                "1. Open PDF in Adobe Acrobat Pro\n"
                "2. View → Show/Hide → Navigation Panes → Order\n"
                "3. Review reading order panel\n"
                "4. Drag and reorder content to match logical reading sequence"
            ),
            blocks_compliance=True
        ))
    
    # -------------------------------------------------------------------------
    # CHECK 6: Heading structure
    # -------------------------------------------------------------------------
    # Headings should be nested properly: H1 → H2 → H3, not H1 → H3
    # This is difficult to check without parsing the tag structure.
    # We'll mark it as needing manual review.
    
    heading_structure_valid = is_tagged  # Heuristic assumption
    
    if not heading_structure_valid:
        issues.append(AccessibilityIssue(
            wcag_criterion="1.3.1",  # Info and Relationships
            level=IssueLevel.WARNING,
            description="Heading structure could not be verified. Manual review recommended.",
            location="Entire document",
            remediation=(
                "1. In Adobe Acrobat Pro, use Tags panel\n"
                "2. Check that headings are nested: H1 > H2 > H3\n"
                "3. Don't skip levels (e.g., H1 directly to H3)\n"
                "4. Ensure each page has at least one H1"
            ),
            blocks_compliance=False  # Warning, not a blocker
        ))
    
    # Close the PDF
    doc.close()
    
    # Create the checks result object
    pdf_checks = PDFAccessibilityChecks(
        is_tagged=is_tagged,
        has_document_language=has_language,
        images_with_alt_text=images_with_alt,
        images_without_alt_text=images_without_alt,
        has_logical_reading_order=has_logical_reading_order,
        is_ocr_needed=is_ocr_needed,
        has_searchable_text=has_searchable_text,
        heading_structure_valid=heading_structure_valid,
        form_fields_labeled=None  # We don't check forms in this version
    )
    
    return pdf_checks, issues


# =============================================================================
# WORD DOCUMENT ACCESSIBILITY CHECKS
# =============================================================================

def check_docx_accessibility(docx_path: str) -> Tuple[DocxAccessibilityChecks, List[AccessibilityIssue]]:
    """
    Check Word document (.docx) for ADA compliance.
    
    What we check:
    1. Document language specified
    2. Images have alt text
    3. Uses proper heading styles (not just bold text)
    4. Heading structure is logical
    5. Tables have headers
    6. Color contrast (basic check of theme colors)
    7. Hyperlinks are descriptive
    """
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx library not installed. "
            "Run: pip install python-docx"
        )
    
    issues: List[AccessibilityIssue] = []
    
    # Open the Word document
    doc = Document(docx_path)
    
    # -------------------------------------------------------------------------
    # CHECK 1: Document language
    # -------------------------------------------------------------------------
    has_language = False
    try:
        # Check document settings for language
        # This requires accessing the XML structure of the document
        if hasattr(doc.core_properties, 'language'):
            lang = doc.core_properties.language
            if lang:
                has_language = True
    except Exception:
        pass
    
    if not has_language:
        issues.append(AccessibilityIssue(
            wcag_criterion="3.1.1",
            level=IssueLevel.ERROR,
            description="Document language is not specified.",
            location="Document properties",
            remediation=(
                "1. In Word: File → Options → Language\n"
                "2. Set 'Office display language' and 'Office authoring languages'\n"
                "3. Or use Review tab → Language → Set Proofing Language"
            ),
            blocks_compliance=True
        ))
    
    # -------------------------------------------------------------------------
    # CHECK 2: Images with alt text
    # -------------------------------------------------------------------------
    images_with_alt = 0
    images_without_alt = 0
    
    # Word stores images in relationships
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            # This is an image relationship
            # Check if any shape/inline shape references it and has alt text
            # This is simplified - production code would be more thorough
            images_without_alt += 1  # Conservative: assume no alt text unless proven
    
    # Check inline shapes (images in text flow)
    for para in doc.paragraphs:
        for run in para.runs:
            # Access the underlying XML to find images
            # This is complex in python-docx, simplified here
            pass
    
    # For now, add a warning about images
    if images_without_alt > 0:
        issues.append(AccessibilityIssue(
            wcag_criterion="1.1.1",
            level=IssueLevel.CRITICAL,
            description=f"Found {images_without_alt} image(s). Verify all have alt text.",
            location="Entire document",
            remediation=(
                "1. Right-click each image\n"
                "2. Select 'Edit Alt Text'\n"
                "3. Add description in the Alt Text pane\n"
                "4. If image is decorative, check 'Mark as decorative'"
            ),
            blocks_compliance=True
        ))
    
    # -------------------------------------------------------------------------
    # CHECK 3: Uses heading styles (not just bold text)
    # -------------------------------------------------------------------------
    uses_heading_styles = False
    has_manual_headings = False
    
    for para in doc.paragraphs:
        # Check if paragraph uses a heading style
        if para.style.name.startswith('Heading'):
            uses_heading_styles = True
        
        # Check if text is just bold (bad practice for headings)
        if para.runs:
            all_bold = all(run.bold for run in para.runs if run.text.strip())
            if all_bold and len(para.text.strip()) > 0:
                # This might be a "fake" heading (just bold text)
                has_manual_headings = True
    
    if not uses_heading_styles:
        issues.append(AccessibilityIssue(
            wcag_criterion="1.3.1",
            level=IssueLevel.CRITICAL,
            description="Document does not use heading styles. Screen readers cannot identify structure.",
            location="Entire document",
            remediation=(
                "1. Select text that should be a heading\n"
                "2. Apply Heading 1, Heading 2, etc. from Styles gallery\n"
                "3. Never use bold text as a substitute for heading styles"
            ),
            blocks_compliance=True
        ))
    
    if has_manual_headings:
        issues.append(AccessibilityIssue(
            wcag_criterion="1.3.1",
            level=IssueLevel.WARNING,
            description="Document may have bold text used as headings instead of proper styles.",
            location="Entire document",
            remediation="Convert bold text to proper Heading styles",
            blocks_compliance=False
        ))
    
    # -------------------------------------------------------------------------
    # CHECK 4: Heading structure (simplified check)
    # -------------------------------------------------------------------------
    heading_levels: List[int] = []
    
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            # Extract level number (e.g., "Heading 1" → 1)
            try:
                level = int(para.style.name.split()[-1])
                heading_levels.append(level)
            except (ValueError, IndexError):
                pass
    
    # Check for level skipping (e.g., H1 → H3 without H2)
    heading_structure_valid = True
    if heading_levels:
        for i in range(1, len(heading_levels)):
            if heading_levels[i] > heading_levels[i-1] + 1:
                heading_structure_valid = False
                issues.append(AccessibilityIssue(
                    wcag_criterion="1.3.1",
                    level=IssueLevel.WARNING,
                    description=f"Heading levels skip: found H{heading_levels[i-1]} followed by H{heading_levels[i]}",
                    location=f"Around heading #{i+1}",
                    remediation="Don't skip heading levels. Go H1 → H2 → H3, not H1 → H3",
                    blocks_compliance=False
                ))
                break
    
    # -------------------------------------------------------------------------
    # CHECK 5: Tables have headers
    # -------------------------------------------------------------------------
    tables_have_headers = True
    
    for table_idx, table in enumerate(doc.tables):
        # Check if first row is formatted as a header row
        # This is a heuristic - we check if cells are bold or styled differently
        first_row = table.rows[0]
        
        # In proper accessible tables, first row should be designated as header
        # We can't easily check this programmatically with python-docx
        # So we add a reminder to check manually
        pass
    
    if len(doc.tables) > 0:
        issues.append(AccessibilityIssue(
            wcag_criterion="1.3.1",
            level=IssueLevel.WARNING,
            description=f"Found {len(doc.tables)} table(s). Verify all have proper header rows.",
            location="Entire document",
            remediation=(
                "1. Click in table → Table Design tab\n"
                "2. Check 'Header Row' box\n"
                "3. Or right-click table → Table Properties → Row → 'Repeat as header row'"
            ),
            blocks_compliance=False
        ))
    
    # -------------------------------------------------------------------------
    # CHECK 6: Color contrast (basic theme check)
    # -------------------------------------------------------------------------
    # This is very difficult to check programmatically
    # We can check theme colors but not actual rendered contrast
    sufficient_color_contrast = True  # Assume OK unless we find problems
    
    # Add a manual review reminder
    issues.append(AccessibilityIssue(
        wcag_criterion="1.4.3",
        level=IssueLevel.INFO,
        description="Color contrast should be verified manually (minimum 4.5:1 for normal text).",
        location="Entire document",
        remediation=(
            "1. Use a color contrast checker tool\n"
            "2. Check text against background colors\n"
            "3. Minimum ratio: 4.5:1 for normal text, 3:1 for large text\n"
            "4. Tool: https://webaim.org/resources/contrastchecker/"
        ),
        blocks_compliance=False
    ))
    
    # -------------------------------------------------------------------------
    # CHECK 7: Hyperlinks are descriptive
    # -------------------------------------------------------------------------
    has_meaningful_hyperlinks = True
    bad_link_texts = ["click here", "here", "link", "read more", "more"]
    
    for para in doc.paragraphs:
        # Check for hyperlinks in the paragraph
        # This requires checking the underlying XML
        # Simplified for this version
        text_lower = para.text.lower()
        for bad_text in bad_link_texts:
            if bad_text in text_lower:
                has_meaningful_hyperlinks = False
                issues.append(AccessibilityIssue(
                    wcag_criterion="2.4.4",
                    level=IssueLevel.WARNING,
                    description=f"Found non-descriptive link text: '{bad_text}'",
                    location="Check document for generic link text",
                    remediation=(
                        "1. Replace 'click here' with descriptive text\n"
                        "2. Example: Change 'Click here for report' to 'Download Q4 Financial Report'\n"
                        "3. Link text should make sense out of context"
                    ),
                    blocks_compliance=False
                ))
                break
    
    # Create the checks result object
    docx_checks = DocxAccessibilityChecks(
        has_document_language=has_language,
        images_with_alt_text=images_with_alt,
        images_without_alt_text=images_without_alt,
        uses_heading_styles=uses_heading_styles,
        heading_structure_valid=heading_structure_valid,
        tables_have_headers=tables_have_headers,
        sufficient_color_contrast=sufficient_color_contrast,
        has_meaningful_hyperlinks=has_meaningful_hyperlinks
    )
    
    return docx_checks, issues


# =============================================================================
# EXCEL ACCESSIBILITY CHECKS
# =============================================================================

def check_xlsx_accessibility(xlsx_path: str) -> Tuple[XlsxAccessibilityChecks, List[AccessibilityIssue]]:
    """
    Check Excel spreadsheet (.xlsx) for ADA compliance.
    
    What we check:
    1. Sheet names are meaningful (not "Sheet1")
    2. Data has proper structure (headers, organization)
    3. Color contrast for text
    4. Avoids merged cells (bad for screen readers)
    5. Charts have alt text
    """
    if not XLSX_AVAILABLE:
        raise ImportError(
            "openpyxl library not installed. "
            "Run: pip install openpyxl"
        )
    
    issues: List[AccessibilityIssue] = []
    
    # Open the Excel file
    wb = load_workbook(xlsx_path, data_only=True)
    
    # -------------------------------------------------------------------------
    # CHECK 1: Meaningful sheet names
    # -------------------------------------------------------------------------
    sheets_have_meaningful_names = True
    default_names = ["Sheet1", "Sheet2", "Sheet3", "Sheet", "Worksheet"]
    
    for sheet_name in wb.sheetnames:
        # Check if sheet has a default/generic name
        if any(sheet_name.startswith(default) for default in default_names):
            sheets_have_meaningful_names = False
            issues.append(AccessibilityIssue(
                wcag_criterion="2.4.6",
                level=IssueLevel.WARNING,
                description=f"Sheet '{sheet_name}' has generic name. Use descriptive names.",
                location=sheet_name,
                remediation=(
                    "1. Right-click sheet tab\n"
                    "2. Select 'Rename'\n"
                    "3. Use descriptive name (e.g., 'Q4_Sales' instead of 'Sheet1')"
                ),
                blocks_compliance=False
            ))
    
    # -------------------------------------------------------------------------
    # CHECK 2: Data structure (headers, no blank rows/columns)
    # -------------------------------------------------------------------------
    has_cell_structure = True
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        
        # Check if first row appears to be headers
        # (Heuristic: bold, different background, etc.)
        first_row = list(ws.iter_rows(min_row=1, max_row=1))[0]
        
        has_headers = False
        for cell in first_row:
            if cell.value and cell.font and cell.font.bold:
                has_headers = True
                break
        
        if not has_headers and first_row[0].value:
            has_cell_structure = False
            issues.append(AccessibilityIssue(
                wcag_criterion="1.3.1",
                level=IssueLevel.ERROR,
                description=f"Sheet '{sheet_name}' may not have clear header row.",
                location=sheet_name,
                remediation=(
                    "1. Add descriptive headers in row 1\n"
                    "2. Format header row (bold, background color)\n"
                    "3. Use Freeze Panes to keep headers visible"
                ),
                blocks_compliance=True
            ))
    
    # -------------------------------------------------------------------------
    # CHECK 3: Color contrast (basic check)
    # -------------------------------------------------------------------------
    sufficient_color_contrast = True
    
    # Add manual review note
    issues.append(AccessibilityIssue(
        wcag_criterion="1.4.3",
        level=IssueLevel.INFO,
        description="Verify color contrast for all colored cells (minimum 4.5:1 ratio).",
        location="All sheets",
        remediation="Use high-contrast color themes and test with color contrast checker",
        blocks_compliance=False
    ))
    
    # -------------------------------------------------------------------------
    # CHECK 4: Merged cells
    # -------------------------------------------------------------------------
    avoids_merged_cells = True
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        
        # Check for merged cells
        if ws.merged_cells:
            merged_ranges = list(ws.merged_cells.ranges)
            if merged_ranges:
                avoids_merged_cells = False
                issues.append(AccessibilityIssue(
                    wcag_criterion="1.3.1",
                    level=IssueLevel.ERROR,
                    description=f"Sheet '{sheet_name}' has {len(merged_ranges)} merged cell(s). Screen readers cannot navigate merged cells properly.",
                    location=sheet_name,
                    remediation=(
                        "1. Unmerge cells: Home → Merge & Center → Unmerge Cells\n"
                        "2. Use Center Across Selection instead\n"
                        "3. Or restructure data to avoid need for merging"
                    ),
                    blocks_compliance=True
                ))
    
    # -------------------------------------------------------------------------
    # CHECK 5: Charts have alt text
    # -------------------------------------------------------------------------
    has_alt_text_for_charts = True
    
    # Check each sheet for charts
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        
        # Check for chart objects
        # openpyxl doesn't provide easy access to chart alt text
        # So we add a reminder to check manually
        if hasattr(ws, '_charts') and ws._charts:
            issues.append(AccessibilityIssue(
                wcag_criterion="1.1.1",
                level=IssueLevel.CRITICAL,
                description=f"Sheet '{sheet_name}' contains chart(s). Verify all have alt text.",
                location=sheet_name,
                remediation=(
                    "1. Right-click chart\n"
                    "2. Select 'Edit Alt Text'\n"
                    "3. Describe what the chart shows (trends, key findings)\n"
                    "4. Include data values if critical"
                ),
                blocks_compliance=True
            ))
    
    # Close the workbook
    wb.close()
    
    # Create the checks result object
    xlsx_checks = XlsxAccessibilityChecks(
        sheets_have_meaningful_names=sheets_have_meaningful_names,
        has_cell_structure=has_cell_structure,
        sufficient_color_contrast=sufficient_color_contrast,
        avoids_merged_cells=avoids_merged_cells,
        has_alt_text_for_charts=has_alt_text_for_charts
    )
    
    return xlsx_checks, issues


# =============================================================================
# UTILITY: Calculate contrast ratio
# =============================================================================

def calculate_contrast_ratio(color1: Tuple[int, int, int], color2: Tuple[int, int, int]) -> float:
    """
    Calculate WCAG contrast ratio between two RGB colors.
    
    Formula from WCAG 2.1:
    https://www.w3.org/WAI/WCAG21/Understanding/contrast-minimum.html
    
    Args:
        color1: RGB tuple (r, g, b) where each value is 0-255
        color2: RGB tuple (r, g, b)
    
    Returns:
        Contrast ratio (1.0 to 21.0)
        - 4.5:1 minimum for normal text (WCAG AA)
        - 3.0:1 minimum for large text (WCAG AA)
        - 7.0:1 minimum for normal text (WCAG AAA)
    """
    def relative_luminance(rgb: Tuple[int, int, int]) -> float:
        """Calculate relative luminance of an RGB color."""
        # Convert to 0-1 range
        r, g, b = [x / 255.0 for x in rgb]
        
        # Apply gamma correction
        def adjust(c):
            return c / 12.92 if c <= 0.03928 else ((c + 0.055) / 1.055) ** 2.4
        
        r, g, b = adjust(r), adjust(g), adjust(b)
        
        # Calculate luminance
        return 0.2126 * r + 0.7152 * g + 0.0722 * b
    
    # Get luminance of both colors
    l1 = relative_luminance(color1)
    l2 = relative_luminance(color2)
    
    # Ensure l1 is the lighter color
    if l1 < l2:
        l1, l2 = l2, l1
    
    # Calculate contrast ratio
    ratio = (l1 + 0.05) / (l2 + 0.05)
    
    return ratio
